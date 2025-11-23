import streamlit as st
import openai as genai
import pdfplumber
import docx
from io import BytesIO
import json
import datetime
import pandas as pd
import time
import base64
from astrapy import DataAPIClient

# --- CONFIGURATION ---
st.set_page_config(
    page_title="AI Resume Architect",
    layout="wide",
    page_icon="🚀",
    initial_sidebar_state="collapsed"
)

# --- 1. DATABASE CONNECTION ---
@st.cache_resource
def get_db_collection():
    """
    Connects to Astra DB with retry logic and configures indexing
    to bypass 8KB limit on large text fields.
    """
    # Load secrets
    try:
        token = st.secrets["ASTRA_DB_APPLICATION_TOKEN"]
        endpoint = st.secrets["ASTRA_DB_API_ENDPOINT"]
    except Exception:
        return None

    client = DataAPIClient(token)
    db = client.get_database_by_api_endpoint(endpoint)
    COLLECTION_NAME = "resume_transactions_python_v1"

    # Retry logic for Serverless Cold Starts
    max_retries = 3
    for attempt in range(max_retries):
        try:
            existing_collections = db.list_collection_names()
            
            if COLLECTION_NAME in existing_collections:
                return db.get_collection(COLLECTION_NAME)
            else:
                # Create collection with NO indexing on large text fields
                db.command({
                    "createCollection": {
                        "name": COLLECTION_NAME,
                        "options": {
                            "indexing": {
                                "deny": [
                                    "original_resume_text", 
                                    "generated_resume", 
                                    "generated_cover_letter", 
                                    "job_description",
                                    "original_file_base64"
                                ]
                            }
                        }
                    }
                })
                return db.get_collection(COLLECTION_NAME)

        except Exception as e:
            if "timeout" in str(e).lower() and attempt < max_retries - 1:
                time.sleep(3)
                continue
            return None
    return None

def save_transaction_to_db(data):
    collection = get_db_collection()
    if collection:
        try:
            collection.insert_one(data)
            return True
        except Exception as e:
            st.error(f"DB Save Error: {e}")
            return False
    return False

def fetch_transactions():
    collection = get_db_collection()
    if not collection: return []
    try:
        # Fetch last 50 transactions
        cursor = collection.find({}, sort={"timestamp": -1}, limit=50)
        return list(cursor)
    except Exception:
        return []

# --- 2. HELPER FUNCTIONS ---

def extract_text(uploaded_file):
    text = ""
    try:
        if uploaded_file.type == "application/pdf":
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(uploaded_file)
            text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        st.error(f"Error reading file: {e}")
        return None

def create_docx(text):
    doc = docx.Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def file_to_base64(uploaded_file):
    """Convert uploaded file to base64 string for DB storage"""
    try:
        bytes_data = uploaded_file.getvalue()
        return base64.b64encode(bytes_data).decode('utf-8')
    except:
        return ""

def base64_to_bytes(base64_string):
    """Convert base64 string back to bytes for download"""
    return base64.b64decode(base64_string)

# --- 3. AI SERVICES ---

def get_gemini_model():
    try:
        api_key = st.secrets["GOOGLE_API_KEY"]
        genai.configure(api_key=api_key)
        # Using 1.5 Flash for speed and large context
        return genai.GenerativeModel('gemini-1.5-flash')
    except Exception:
        st.error("Google API Key not found in secrets.")
        return None

def analyze_resume(text, jd):
    model = get_gemini_model()
    if not model: return {"match_score": 0, "tips": []}
    
    prompt = f"""
    Act as a strict ATS. Compare the Resume against the Job Description.
    Output a raw JSON object with these keys: "match_score" (number 0-100), "missing_keywords" (list of strings), "tips" (list of strings).
    Do not output markdown code blocks. Just the JSON.
    
    RESUME: {text[:10000]}
    JD: {jd[:5000]}
    """
    try:
        response = model.generate_content(prompt)
        # Clean response if model adds markdown blocks
        clean_text = response.text.replace("```json", "").replace("```", "")
        return json.loads(clean_text)
    except Exception as e:
        st.error(f"Analysis Error: {e}")
        return {"match_score": 0, "tips": ["Error analyzing resume"]}

def optimize_resume(text, jd):
    model = get_gemini_model()
    if not model: return ""
    
    prompt = f"""
    Rewrite this resume to get a 95% match score against the JD.
    Use "Keyword Mirroring" (use exact phrasing from JD).
    Keep actual companies/dates but rewrite bullets to focus on results.
    Return clean Markdown.
    
    RESUME: {text}
    JD: {jd}
    """
    response = model.generate_content(prompt)
    return response.text

def generate_cover_letter(text, jd, length_type):
    model = get_gemini_model()
    if not model: return ""
    
    length_prompt = {
        "Condensed": "Keep it under 200 words. Punchy and direct.",
        "Medium": "Standard professional length (300 words). Balanced.",
        "Elaborate": "Detailed storytelling (450+ words). Deep dive into achievements."
    }
    
    prompt = f"""
    Write a cover letter based on this resume and JD.
    Style: {length_prompt.get(length_type, "Medium")}
    Return clean Markdown.
    
    RESUME: {text}
    JD: {jd}
    """
    response = model.generate_content(prompt)
    return response.text

# --- 4. UI PAGES ---

def generator_page():
    st.title("🚀 AI Resume Architect")
    
    # Init session state
    if "generated" not in st.session_state:
        st.session_state.generated = None

    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("1. Upload Resume")
        uploaded_file = st.file_uploader("Upload PDF or DOCX", type=["pdf", "docx"])
        
        st.subheader("3. Options")
        cl_length = st.select_slider("Cover Letter Length", options=["Condensed", "Medium", "Elaborate"], value="Medium")

    with col2:
        st.subheader("2. Job Description")
        jd_text = st.text_area("Paste JD here", height=200)

    if st.button("Generate Application", type="primary"):
        if not uploaded_file or not jd_text:
            st.warning("Please provide both a resume and a job description.")
            return
        
        with st.status("Architecting Application...", expanded=True) as status:
            # 1. Extract Text
            status.write("Reading document...")
            resume_text = extract_text(uploaded_file)
            
            if not resume_text:
                status.update(label="Failed to read file", state="error")
                return

            # 2. Analyze Original
            status.write("Analyzing original match score...")
            original_analysis = analyze_resume(resume_text, jd_text)
            
            # 3. Optimize Resume & Cover Letter
            status.write("Optimizing resume & drafting cover letter...")
            optimized_text = optimize_resume(resume_text, jd_text)
            cover_letter_text = generate_cover_letter(resume_text, jd_text, cl_length)
            
            # 4. Analyze New
            status.write("Verifying new match score...")
            new_analysis = analyze_resume(optimized_text, jd_text)
            
            # 5. Save to DB
            status.write("Saving transaction to Astra DB...")
            timestamp = datetime.datetime.now().isoformat()
            
            transaction_data = {
                "timestamp": timestamp,
                "job_title": jd_text.split('\n')[0][:50],
                "job_description": jd_text,
                "original_filename": uploaded_file.name,
                "original_file_base64": file_to_base64(uploaded_file),
                "original_score": original_analysis.get('match_score', 0),
                "optimized_score": new_analysis.get('match_score', 0),
                "critical_keywords": new_analysis.get('missing_keywords', []), # Storing new analysis keywords as 'targeted'
                "improvements": original_analysis.get('tips', []),
                "original_resume_text": resume_text,
                "generated_resume": optimized_text,
                "generated_cover_letter": cover_letter_text
            }
            
            save_transaction_to_db(transaction_data)
            
            # Save to session state to prevent reload loss
            st.session_state.generated = {
                "original_stats": original_analysis,
                "new_stats": new_analysis,
                "optimized_resume": optimized_text,
                "cover_letter": cover_letter_text
            }
            
            status.update(label="Complete!", state="complete", expanded=False)

    # --- RESULTS DISPLAY ---
    if st.session_state.generated:
        res = st.session_state.generated
        
        st.divider()
        
        # Scoreboard
        c1, c2, c3 = st.columns(3)
        with c1:
            st.metric("Original Score", f"{res['original_stats'].get('match_score', 0)}%")
        with c2:
            st.metric("Optimized Score", f"{res['new_stats'].get('match_score', 0)}%", 
                     delta=res['new_stats'].get('match_score', 0) - res['original_stats'].get('match_score', 0))
        with c3:
            st.info(f"**Top Tip:** {res['original_stats'].get('tips', [''])[0]}")

        # Downloads & Previews
        tab1, tab2 = st.tabs(["📄 Optimized Resume", "✉️ Cover Letter"])
        
        with tab1:
            col_d1, col_d2 = st.columns([1, 4])
            with col_d1:
                st.download_button(
                    "Download .docx",
                    data=create_docx(res['optimized_resume']),
                    file_name="Optimized_Resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            with col_d2:
                st.markdown(res['optimized_resume'])
                
        with tab2:
            col_d1, col_d2 = st.columns([1, 4])
            with col_d1:
                st.download_button(
                    "Download .docx",
                    data=create_docx(res['cover_letter']),
                    file_name="Cover_Letter.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            with col_d2:
                st.markdown(res['cover_letter'])

def admin_page():
    st.header("🔒 Admin Dashboard")
    
    # Simple password check against secrets
    try:
        admin_pass = st.secrets["ADMIN_PASSWORD"]
    except:
        st.error("ADMIN_PASSWORD not set in secrets.")
        return

    pwd_input = st.text_input("Enter Admin Password", type="password")
    
    if pwd_input == admin_pass:
        if st.button("Refresh Data"):
            st.rerun()
            
        transactions = fetch_transactions()
        
        if not transactions:
            st.info("No transactions found in DB.")
            return
            
        # Summary Table
        df = pd.DataFrame(transactions)
        st.dataframe(
            df[['timestamp', 'job_title', 'original_score', 'optimized_score', 'original_filename']],
            use_container_width=True,
            hide_index=True
        )
        
        st.divider()
        
        # Detailed View
        tx_options = {f"{t['timestamp']} - {t['job_title']}": t for t in transactions}
        selected_tx_key = st.selectbox("Select Transaction Details", list(tx_options.keys()))
        
        if selected_tx_key:
            tx = tx_options[selected_tx_key]
            
            c1, c2, c3 = st.columns(3)
            c1.metric("Original", f"{tx.get('original_score')}%")
            c2.metric("Optimized", f"{tx.get('optimized_score')}%")
            
            # Downloads
            d1, d2, d3 = st.columns(3)
            
            # Original File Download
            if tx.get('original_file_base64'):
                try:
                    file_bytes = base64_to_bytes(tx['original_file_base64'])
                    d1.download_button("⬇️ Original File", file_bytes, tx.get('original_filename', 'resume'))
                except:
                    d1.error("File corrupted")
            
            d2.download_button("⬇️ Optimized Resume", create_docx(tx.get('generated_resume', '')), "Optimized.docx")
            d3.download_button("⬇️ Cover Letter", create_docx(tx.get('generated_cover_letter', '')), "CoverLetter.docx")

            with st.expander("View Job Description"):
                st.text(tx.get('job_description'))
                
            with st.expander("View Improvements Made"):
                st.write(tx.get('improvements'))

# --- 5. MAIN APP ROUTER ---

def main():
    # Sidebar Navigation
    st.sidebar.title("Navigation")
    page = st.sidebar.radio("Go to", ["Generator", "Admin Dashboard"])
    
    if page == "Generator":
        generator_page()
    else:
        admin_page()

if __name__ == "__main__":
    main()

